"""
Tool: DOCX / DOC text extractor.

Handles both modern .docx (python-docx) and legacy .doc (LibreOffice conversion).
Preserves paragraph order and extracts text from tables too
(many CVs put skills/experience in tables).
"""
from __future__ import annotations

import logging
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path

logger = logging.getLogger(__name__)


@dataclass
class ExtractionResult:
    text: str
    method: str       # "docx" | "doc_converted"
    paragraph_count: int
    char_count: int


def extract_docx_text(file_path: str) -> ExtractionResult:
    """
    Extract all text from a .docx or .doc file.
    Includes body paragraphs AND table cell text.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"CV file not found: {file_path}")

    suffix = path.suffix.lower()

    if suffix == ".docx":
        return _extract_docx(path)
    elif suffix == ".doc":
        # Legacy .doc — convert to .docx via LibreOffice first
        return _extract_legacy_doc(path)
    else:
        raise ValueError(f"Expected .docx or .doc, got: {suffix}")


def _extract_docx(path: Path) -> ExtractionResult:
    """Extract from modern .docx using python-docx."""
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []

    # --- Body paragraphs ---
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # --- Tables (CV skills/experience often in table cells) ---
    for table in doc.tables:
        for row in table.rows:
            row_cells: list[str] = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_cells.append(cell_text)
            if row_cells:
                parts.append(" | ".join(row_cells))

    full_text = "\n".join(parts)
    logger.info(
        "DOCX extracted: %s (%d chars, %d paragraphs)",
        path.name, len(full_text), len(parts),
    )

    return ExtractionResult(
        text=full_text,
        method="docx",
        paragraph_count=len(parts),
        char_count=len(full_text),
    )


def _extract_legacy_doc(path: Path) -> ExtractionResult:
    """
    Convert .doc → .docx via LibreOffice headless, then extract.
    Requires LibreOffice installed: apt install libreoffice
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            subprocess.run(
                [
                    "libreoffice", "--headless", "--convert-to", "docx",
                    "--outdir", tmpdir, str(path),
                ],
                check=True,
                capture_output=True,
                timeout=30,
            )
        except (subprocess.CalledProcessError, FileNotFoundError) as e:
            raise RuntimeError(
                f"LibreOffice not available to convert .doc file: {path.name}. "
                "Install with: apt install libreoffice"
            ) from e

        converted = Path(tmpdir) / path.with_suffix(".docx").name
        if not converted.exists():
            raise RuntimeError(f"LibreOffice conversion produced no output for {path.name}")

        result = _extract_docx(converted)
        return ExtractionResult(
            text=result.text,
            method="doc_converted",
            paragraph_count=result.paragraph_count,
            char_count=result.char_count,
        )